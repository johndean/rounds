"""
Artifact transformer — produces docx / srt / vtt / txt / zip from a
session's segments + slides + speakers + normalization.

Ports MIC `app/engines/artifact_transformer.py` (540 LOC). Each public
function returns raw bytes for the caller to stream out via FastAPI.

Phase 6p / U141-U142. Closes audit gap 🟠 #11.

Related ADRs: ADR-004 (single-source export engine).
Related business rules: BR-016 (format-specific markup stripping), BR-017 (Unknown speaker fallback).
"""
from __future__ import annotations

import io
import json
import logging
import re
import zipfile
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class SegmentForExport:
    seq: int
    start_ms: int
    end_ms: int
    text: str
    slide_index: int | None
    slide_title: str | None
    speaker_name: str | None


@dataclass
class SlideForExport:
    slide_index: int
    title: str
    full_text: str
    bullets: list[str]


@dataclass
class PollForExport:
    slide_index: int
    question: str
    options: list[dict]  # [{label, count, percent}]


@dataclass
class ChatForExport:
    author: str
    body: str
    sent_at_ms: int


@dataclass
class SessionForExport:
    code: str
    title: str
    presenter: str | None
    duration_sec: int | None
    segments: list[SegmentForExport]
    slides: list[SlideForExport]
    polls: list[PollForExport] = field(default_factory=list)
    chat: list[ChatForExport] = field(default_factory=list)
    publishing_links: dict = field(default_factory=dict)
    resources: list[dict] = field(default_factory=list)  # [{slide_number, label, url}]


# ─── Helpers ────────────────────────────────────────────────────────────


def _fmt_srt_time(ms: int) -> str:
    hours = ms // 3_600_000
    ms %= 3_600_000
    mins = ms // 60_000
    ms %= 60_000
    secs = ms // 1000
    millis = ms % 1000
    return f"{hours:02d}:{mins:02d}:{secs:02d},{millis:03d}"


def _fmt_vtt_time(ms: int) -> str:
    hours = ms // 3_600_000
    ms %= 3_600_000
    mins = ms // 60_000
    ms %= 60_000
    secs = ms // 1000
    millis = ms % 1000
    return f"{hours:02d}:{mins:02d}:{secs:02d}.{millis:03d}"


# ─── Public exporters ───────────────────────────────────────────────────


def to_txt(session: SessionForExport) -> bytes:
    lines: list[str] = [
        f"# {session.title}".strip(),
        f"Code: {session.code}",
    ]
    if session.presenter:
        lines.append(f"Presenter: {session.presenter}")
    lines.append("")
    last_slide = None
    for seg in session.segments:
        if seg.slide_index != last_slide and seg.slide_title:
            lines.append("")
            lines.append(f"## Slide {seg.slide_index}: {seg.slide_title}")
            lines.append("")
            last_slide = seg.slide_index
        if seg.speaker_name:
            lines.append(f"{seg.speaker_name}: {seg.text}")
        else:
            lines.append(seg.text)
    return ("\n".join(lines) + "\n").encode("utf-8")


def to_srt(session: SessionForExport) -> bytes:
    """
    SRT output. Per-segment text is passed through apply_srt_transform to
    strip any leftover slide markers / speaker labels / [pq] tags / curly
    annotations so captions render as plain speech.
    """
    chunks: list[str] = []
    for i, seg in enumerate(session.segments, start=1):
        chunks.append(str(i))
        chunks.append(f"{_fmt_srt_time(seg.start_ms)} --> {_fmt_srt_time(seg.end_ms)}")
        cleaned = apply_srt_transform(seg.text or "").strip()
        chunks.append(cleaned)
        chunks.append("")
    return "\n".join(chunks).encode("utf-8")


def to_vtt(session: SessionForExport) -> bytes:
    chunks: list[str] = ["WEBVTT", ""]
    for seg in session.segments:
        chunks.append(f"{_fmt_vtt_time(seg.start_ms)} --> {_fmt_vtt_time(seg.end_ms)}")
        chunks.append((seg.text or "").strip())
        chunks.append("")
    return "\n".join(chunks).encode("utf-8")


def to_docx(session: SessionForExport) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(session.title or session.code, level=1)
    if session.presenter:
        doc.add_paragraph(f"Presenter: {session.presenter}")
    doc.add_paragraph(f"Code: {session.code}")
    doc.add_paragraph()

    last_slide = None
    for seg in session.segments:
        if seg.slide_index != last_slide and seg.slide_title:
            doc.add_heading(f"Slide {seg.slide_index}: {seg.slide_title}", level=2)
            last_slide = seg.slide_index
        # Phase 5 (2026-06-05) — preserve segment formatting in the
        # DOCX export. Previously every segment collapsed to a single
        # paragraph regardless of embedded \n, so an editor who hard-
        # wrapped a long quote into 3 paragraphs would see one wall of
        # text in the exported Word doc.
        # Convention: '\n\n' (hard return) = new paragraph,
        #             '\n'   (soft return) = line break within paragraph
        # (matches what MS Word emits for Enter vs Shift+Enter).
        seg_text = seg.text or ""
        paragraphs = seg_text.split("\n\n")
        for p_idx, para_text in enumerate(paragraphs):
            para = doc.add_paragraph()
            if p_idx == 0 and seg.speaker_name:
                speaker_run = para.add_run(f"{seg.speaker_name}: ")
                speaker_run.bold = True
            lines = para_text.split("\n")
            for l_idx, line in enumerate(lines):
                if l_idx > 0:
                    # Soft line break — python-docx exposes it via
                    # add_break() on a run (no arg = WD_BREAK.LINE).
                    para.add_run().add_break()
                if line:
                    para.add_run(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Macro layer (CMS / SRT cleanup) ────────────────────────────────────


def _build_marked_transcript(session: SessionForExport) -> str:
    """
    Build the marked-up source text that macros operate on:
      ++N*+        slide marker
      **Name:**    speaker label
      [pq][t]      placeholder for chat/poll injection at timestamp t

    This is the artifact CMS macros expect as input.
    """
    lines: list[str] = []
    last_slide = None
    last_speaker = None
    for seg in session.segments:
        if seg.slide_index is not None and seg.slide_index != last_slide:
            lines.append("")
            lines.append(f"++{seg.slide_index + 1}*+")
            lines.append("")
            last_slide = seg.slide_index
        if seg.speaker_name and seg.speaker_name != last_speaker:
            lines.append(f"**{seg.speaker_name}:**")
            last_speaker = seg.speaker_name
        lines.append(seg.text or "")
    return "\n".join(lines).strip()


def apply_srt_transform(text: str) -> str:
    """
    11-step deterministic SRT macro — strips structural markup leaving only
    speech-as-text. Verbatim port of MIC artifact_transformer._apply_srt_transform.
    """
    # BR-016 — Format-specific markup stripping rule.
    # See docs/BUSINESS_RULES.md#br-016.
    # Why: to_srt() calls this transform to produce clean caption text
    # (slide codes / speaker labels / [pq] tags / curly annotations removed).
    # to_vtt() deliberately does NOT call it — VTT captions preserve
    # structural markup so editors can correlate captions to anchors.
    # to_docx() / to_txt() apply their own format-specific cleanup. Filler
    # words ("um"/"uh"/...) are stripped earlier at the normalize phase
    # (app/iil/normalization.py:TIER1_WORDS), not here.
    t = text
    # 1. Slide codes
    t = re.sub(r"\+\+\d+\*\+\s*", "", t)
    # 2. [Video] tags
    t = re.sub(r"\[\s*[Vv]ideo\s*\]", "", t)
    # 3. Speaker labels
    t = re.sub(r"<b[^>]*>[^<]+:</b>\s*", "", t)
    t = re.sub(r"\*\*[^*]+:\*\*\s*", "", t)
    t = re.sub(r"^[A-Z][a-zA-Z\s]+:\s+", "", t, flags=re.MULTILINE)
    # 4. [pq][HH:MM:SS] timestamps
    t = re.sub(r"\[pq\]\[\d{1,2}:\d{2}(?::\d{2})?\]\s*", "", t)
    # 5. Bare [pq]
    t = re.sub(r"\[pq\]\s*", "", t)
    # 6. {curly} — keep contents
    t = re.sub(r"\{([^}]*)\}", r"\1", t)
    # 7. Poll markers
    t = re.sub(r"\[Poll\s*#?\d+\]", "", t, flags=re.IGNORECASE)
    t = re.sub(r"Poll\s+#?\d+\s*\n", "", t, flags=re.IGNORECASE)
    # 8. Double-space collapse
    t = re.sub(r"  +", " ", t)
    # 9. Strip per-line
    t = "\n".join(line.strip() for line in t.split("\n"))
    # 10. Triple+ newlines → double
    t = re.sub(r"\n{3,}", "\n\n", t)
    # 11. Final strip
    return t.strip()


class CMSValidationError(Exception):
    """CMS doc validation gate failed — unresolved markers detected."""


def apply_cms_transform(
    text: str,
    polls: list[PollForExport],
    chat: list[ChatForExport],
    resources: list[dict],
    *,
    hyperlinks: dict[str, str] | None = None,
    strict: bool = False,
) -> str:
    """
    CMS macro (#65 + #66 + #69) — 9-step publish-ready transform:
      1. Strip {curly} content (CMS variant — removes, doesn't keep)
      2. Whitespace normalize
      3. Inject poll blocks at slide markers
      4. Replace [pq][HH:MM:SS] with nearest chat at that timestamp (#69)
      5. Strip leftover bare [pq] tokens
      6. Replace inline {{token}} hyperlinks per `hyperlinks` map (#65 step-7)
      7. Append Resources section
      8. Final whitespace cleanup
      9. Validate — reject if any unresolved [X][T=]/curly/slide-marker remain (#66)
    """
    t = text
    # 1. Strip curly text (CMS variant — content removed)
    t = re.sub(r"\{[^}]*\}", "", t)
    # 2. Whitespace normalize
    t = "\n".join(line.strip() for line in t.split("\n"))
    t = re.sub(r"  +", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)

    # 3. Inject polls after their slide markers
    for poll in polls:
        marker = f"++{poll.slide_index + 1}*+"
        block = _format_poll_block(poll)
        t = t.replace(marker, f"{marker}\n\n{block}", 1)

    # 4. Replace [pq][HH:MM:SS] markers with the nearest chat at that timestamp (#69)
    chat_by_ms = sorted(chat, key=lambda c: c.sent_at_ms or 0)

    def _replace_pq(match: re.Match) -> str:
        ts_str = match.group(1)
        target_ms = _parse_hhmmss_to_ms(ts_str)
        if target_ms is None or not chat_by_ms:
            return ""
        nearest = min(chat_by_ms, key=lambda c: abs((c.sent_at_ms or 0) - target_ms))
        ts = _fmt_vtt_time(nearest.sent_at_ms).split(".")[0]
        return f"**{nearest.author}** ({ts}): {nearest.body}\n\n"

    t = re.sub(r"\[pq\]\[(\d{1,2}:\d{2}(?::\d{2})?)\]\s*", _replace_pq, t)

    # 5. Strip any leftover bare [pq] tokens.
    t = re.sub(r"\[pq\]\s*", "", t)

    # 6. Hyperlink replacement (#65 step-7) — replace {{token}} per hyperlinks dict.
    if hyperlinks:
        for token, url in hyperlinks.items():
            pattern = r"\{\{\s*" + re.escape(token) + r"\s*\}\}"
            t = re.sub(pattern, f"[{token}]({url})", t)

    # 7. Resources section
    if resources:
        t = t.rstrip() + "\n\n---\n\n**Resources**\n\n"
        for r in resources:
            label = r.get("label") or r.get("url") or ""
            url = r.get("url") or ""
            slide_n = r.get("slide_number")
            prefix = f"(Slide {slide_n}) " if slide_n else ""
            if url:
                t += f"- {prefix}[{label}]({url})\n"
            else:
                t += f"- {prefix}{label}\n"

    # 8. Final whitespace cleanup
    t = re.sub(r"\n{3,}", "\n\n", t).strip()

    # 9. Validate — reject unresolved markers (#66)
    if strict:
        _validate_cms_doc(t)

    return t


def _parse_hhmmss_to_ms(ts: str) -> Optional[int]:
    """Parse 'HH:MM:SS' or 'MM:SS' to ms. Returns None on malformed input."""
    parts = ts.split(":")
    try:
        if len(parts) == 3:
            h, m, s = int(parts[0]), int(parts[1]), int(parts[2])
            return (h * 3600 + m * 60 + s) * 1000
        if len(parts) == 2:
            m, s = int(parts[0]), int(parts[1])
            return (m * 60 + s) * 1000
    except ValueError:
        return None
    return None


_UNRESOLVED_PATTERNS = [
    re.compile(r"\[X\]"),                         # editor placeholder
    re.compile(r"\[T=[^\]]*\]"),                  # unresolved timestamp token
    re.compile(r"\{[^}]+\}"),                     # leftover curly
    re.compile(r"\{\{[^}]+\}\}"),                 # unreplaced hyperlink token
    re.compile(r"\[pq\](?:\[[^\]]*\])?"),         # any leftover pq marker
]


def _validate_cms_doc(text: str) -> None:
    """
    Reject CMS doc that still contains unresolved medical-review markers
    or formatting placeholders. Closes 🟠 #66.
    """
    failures: list[str] = []
    for pattern in _UNRESOLVED_PATTERNS:
        match = pattern.search(text)
        if match:
            failures.append(f"unresolved marker `{match.group(0)}` at index {match.start()}")
    if failures:
        raise CMSValidationError("; ".join(failures))


# ─── Caption-line validator ─────────────────────────────────────────────────


def validate_final_srt(srt_bytes: bytes, *, max_line_chars: int = 42) -> None:
    """
    DCMP caption compliance check (#67) — line length ≤42 chars, no HTML, no
    markers, no curly braces. Raises CMSValidationError on violation.
    """
    text = srt_bytes.decode("utf-8") if isinstance(srt_bytes, (bytes, bytearray)) else str(srt_bytes)
    failures: list[str] = []
    for i, line in enumerate(text.splitlines(), start=1):
        if re.match(r"^\d+$", line.strip()):
            continue  # cue number
        if "-->" in line:
            continue  # timestamp line
        if "<" in line and ">" in line:
            failures.append(f"line {i}: contains HTML tag")
        if "{" in line or "}" in line:
            failures.append(f"line {i}: contains curly marker")
        if "[pq]" in line.lower() or "[x]" in line.lower() or "[t=" in line.lower():
            failures.append(f"line {i}: contains unresolved marker")
        stripped = line.strip()
        if stripped and len(stripped) > max_line_chars:
            failures.append(f"line {i}: {len(stripped)} chars exceeds {max_line_chars}")
    if failures:
        raise CMSValidationError("validate_final_srt: " + "; ".join(failures[:8]))


def _format_poll_block(poll: PollForExport) -> str:
    lines = [f"***{poll.question}***"]
    for opt in poll.options:
        count = opt.get("count", 0)
        pct = opt.get("percent", 0)
        label = opt.get("label", "")
        lines.append(f"{count} ({pct}%) {label}")
    return "\n".join(lines)


def to_cms_html(session: SessionForExport) -> bytes:
    """
    Publish-ready CMS output: marked transcript → CMS macro → light HTML.
    Slide markers become <h2>, speaker labels stay as bold prefixes.
    """
    marked = _build_marked_transcript(session)
    # Publish path — enforce validation gate (#66). Unresolved markers raise CMSValidationError.
    cms = apply_cms_transform(
        marked, session.polls, session.chat, session.resources,
        hyperlinks=getattr(session, "hyperlinks", None),
        strict=True,
    )

    # Convert markdown-ish constructs to inline HTML.
    html_lines = ["<!doctype html>",
                  '<html><head><meta charset="utf-8">',
                  f"<title>{_escape(session.title or session.code)}</title>",
                  "<style>",
                  "  body{font-family:Georgia,serif;max-width:780px;margin:32px auto;padding:0 16px;line-height:1.6;color:#222}",
                  "  h1{font-size:28px;margin-bottom:4px}",
                  "  h2{font-size:18px;margin-top:32px;border-bottom:1px solid #ddd;padding-bottom:4px}",
                  "  p{margin:8px 0}",
                  "  strong{font-weight:700}",
                  "  em{font-style:italic}",
                  "  hr{margin:32px 0;border:none;border-top:1px solid #ddd}",
                  "</style></head><body>",
                  f"<h1>{_escape(session.title or session.code)}</h1>"]
    if session.presenter:
        html_lines.append(f"<p><em>Presented by {_escape(session.presenter)}</em></p>")
    html_lines.append("")

    # Transform the CMS markdown body into HTML.
    body_html = _markdown_to_html(cms)
    html_lines.append(body_html)
    html_lines.append("</body></html>")
    return "\n".join(html_lines).encode("utf-8")


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


_SLIDE_MARKER_RE = re.compile(r"\+\+(\d+)\*\+")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_TRIPLE_BOLD_RE = re.compile(r"\*\*\*(.+?)\*\*\*")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HR_RE = re.compile(r"^---\s*$", re.MULTILINE)


def _markdown_to_html(text: str) -> str:
    """Minimal markdown→HTML for CMS output. Not a full Markdown impl."""
    out = []
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Slide marker → h2
        slide_match = _SLIDE_MARKER_RE.search(block)
        if slide_match and block.strip() == slide_match.group(0):
            out.append(f"<h2>Slide {slide_match.group(1)}</h2>")
            continue
        # Horizontal rule
        if _HR_RE.match(block):
            out.append("<hr/>")
            continue
        # Inline transforms within the block
        block = _TRIPLE_BOLD_RE.sub(r"<strong><em>\1</em></strong>", block)
        block = _BOLD_RE.sub(r"<strong>\1</strong>", block)
        block = _LINK_RE.sub(r'<a href="\2">\1</a>', block)
        # Convert line breaks within a block to <br>
        lines = block.split("\n")
        if all(line.startswith("- ") for line in lines):
            items = "".join(f"<li>{ln[2:].strip()}</li>" for ln in lines)
            out.append(f"<ul>{items}</ul>")
        else:
            inner = "<br/>".join(lines)
            out.append(f"<p>{inner}</p>")
    return "\n".join(out)


def to_zip(session: SessionForExport) -> bytes:
    """Bundle docx + srt + vtt + txt + html + slide bullets into a single zip."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{session.code}.txt",  to_txt(session))
        zf.writestr(f"{session.code}.srt",  to_srt(session))
        zf.writestr(f"{session.code}.vtt",  to_vtt(session))
        zf.writestr(f"{session.code}.docx", to_docx(session))
        zf.writestr(f"{session.code}.html", to_cms_html(session))  # 7e — CMS publish-ready
        slide_lines: list[str] = [f"# {session.title} — slide outline", ""]
        for s in session.slides:
            slide_lines.append(f"## Slide {s.slide_index + 1}: {s.title}")
            for bullet in s.bullets:
                slide_lines.append(f"- {bullet}")
            slide_lines.append("")
        zf.writestr(f"{session.code}_slides.txt", "\n".join(slide_lines).encode("utf-8"))
    return buf.getvalue()


# ─── Data loader ────────────────────────────────────────────────────────


def load_session_for_export(session_id: str) -> SessionForExport:
    """Fetch everything a session export needs in a single read."""
    from sqlalchemy import create_engine, text

    from app.config import settings

    sync_url = settings.DATABASE_URL.replace("+asyncpg", "")
    engine = create_engine(sync_url)
    try:
        with engine.connect() as conn:
            sess = conn.execute(
                text(
                    """
                    SELECT code, title, presenter, duration_sec,
                           coalesce(publishing_links, '{}'::jsonb)
                      FROM sessions WHERE id = CAST(:sid AS uuid)
                    """
                ),
                {"sid": session_id},
            ).fetchone()
            if not sess:
                raise RuntimeError(f"export: session {session_id} not found")

            segments = conn.execute(
                text(
                    """
                    SELECT seg.seq, seg.start_ms, seg.end_ms, seg.text,
                           sl.slide_index, sl.title,
                           sp.name AS speaker_name
                      FROM segments seg
                      LEFT JOIN slides sl   ON sl.id = seg.slide_id
                      LEFT JOIN speakers sp ON sp.id = seg.speaker_id
                     WHERE seg.session_id = CAST(:sid AS uuid)
                     ORDER BY seg.seq ASC
                    """
                ),
                {"sid": session_id},
            ).fetchall()

            slides_rows = conn.execute(
                text(
                    """
                    SELECT sl.slide_index, sl.title, sl.full_text,
                           coalesce(array_agg(b.text ORDER BY b.position) FILTER (WHERE b.text IS NOT NULL), '{}')
                      FROM slides sl
                      LEFT JOIN bullets b ON b.slide_id = sl.id
                     WHERE sl.session_id = CAST(:sid AS uuid)
                     GROUP BY sl.slide_index, sl.title, sl.full_text
                     ORDER BY sl.slide_index ASC
                    """
                ),
                {"sid": session_id},
            ).fetchall()

            # Polls — parsed from manifest (sessions.polls_parsed JSONB).
            polls_row = conn.execute(
                text("SELECT polls_parsed FROM sessions WHERE id = CAST(:sid AS uuid)"),
                {"sid": session_id},
            ).fetchone()

            chat_rows = conn.execute(
                text(
                    """
                    SELECT author, body, sent_at_ms FROM chat_messages
                     WHERE session_id = CAST(:sid AS uuid)
                     ORDER BY sent_at_ms ASC
                    """
                ),
                {"sid": session_id},
            ).fetchall()

            resource_rows = conn.execute(
                text(
                    """
                    SELECT slide_number, label, url FROM session_slide_resources
                     WHERE session_id = CAST(:sid AS uuid)
                     ORDER BY slide_number, sort_order
                    """
                ),
                {"sid": session_id},
            ).fetchall()
    finally:
        engine.dispose()

    parsed_polls = polls_row[0] if polls_row and polls_row[0] else []
    if isinstance(parsed_polls, str):
        try:
            parsed_polls = json.loads(parsed_polls)
        except json.JSONDecodeError:
            parsed_polls = []

    publishing_links = sess[4] if len(sess) > 4 and sess[4] else {}
    if isinstance(publishing_links, str):
        try:
            publishing_links = json.loads(publishing_links)
        except json.JSONDecodeError:
            publishing_links = {}

    return SessionForExport(
        code=sess[0],
        title=sess[1] or sess[0],
        presenter=sess[2],
        duration_sec=sess[3],
        publishing_links=publishing_links,
        polls=[
            PollForExport(
                slide_index=(p.get("slide_n") or 1) - 1,
                question=p.get("question", ""),
                options=p.get("options", []),
            )
            for p in (parsed_polls or [])
        ],
        chat=[
            ChatForExport(
                author=r[0], body=r[1], sent_at_ms=r[2] or 0,
            )
            for r in chat_rows
        ],
        resources=[
            {"slide_number": r[0], "label": r[1], "url": r[2]}
            for r in resource_rows
        ],
        segments=[
            SegmentForExport(
                seq=r[0], start_ms=r[1] or 0, end_ms=r[2] or 0,
                text=r[3] or "", slide_index=r[4],
                slide_title=r[5], speaker_name=r[6],
            )
            for r in segments
        ],
        slides=[
            SlideForExport(
                slide_index=r[0], title=r[1] or f"Slide {r[0] + 1}",
                full_text=r[2] or "", bullets=list(r[3] or []),
            )
            for r in slides_rows
        ],
    )
