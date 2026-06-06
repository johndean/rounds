"""
Phase 5 — speaker.role threaded through DOCX export.

Plan ref: docs/plans/2026-06-05-010-zero-gap-parity-plan.md §Phase 5.
Audit IDs closed: D2 (speaker formatting), E14 (blue/bold speaker AI
auto-format), TR3 (speaker.role schema unused).
"""
from __future__ import annotations


def test_segment_for_export_has_speaker_role_field():
    """Dataclass must expose speaker_role as Optional[str], default None."""
    import dataclasses

    from app.engines.artifact_transformer import SegmentForExport

    fields = {f.name: f for f in dataclasses.fields(SegmentForExport)}
    assert "speaker_role" in fields, "speaker_role missing from SegmentForExport"
    role_field = fields["speaker_role"]
    # default must be None so older callers don't break.
    assert role_field.default is None, "speaker_role must default to None"


def test_load_session_for_export_selects_speaker_role():
    """Source-level: the SQL in load_session_for_export must SELECT sp.role.
    Without this, the column stays unused (TR3 regression)."""
    import pathlib

    src = pathlib.Path(__file__).resolve().parent.parent / "app" / "engines" / "artifact_transformer.py"
    text = src.read_text(encoding="utf-8")
    assert "sp.role" in text, "load_session_for_export does not SELECT sp.role"
    assert "speaker_role=r[7]" in text, "constructor does not pass speaker_role through"


def test_to_docx_applies_navy_for_primary_speaker():
    """to_docx must color the speaker run navy when speaker_role == 'primary'.
    Bold-only fallback when role is None / moderator / guest."""
    from app.engines.artifact_transformer import (
        SegmentForExport,
        SessionForExport,
        to_docx,
    )

    seg_primary = SegmentForExport(
        seq=0, start_ms=0, end_ms=1000, text="hello",
        slide_index=0, slide_title="Slide 1",
        speaker_name="Dr. Mueller", speaker_role="primary",
    )
    seg_moderator = SegmentForExport(
        seq=1, start_ms=1000, end_ms=2000, text="hi",
        slide_index=0, slide_title="Slide 1",
        speaker_name="J. Hsu", speaker_role="moderator",
    )
    session = SessionForExport(
        code="test-1", title="Test session", presenter=None,
        duration_sec=None,
        segments=[seg_primary, seg_moderator],
        slides=[],
    )

    blob = to_docx(session)
    assert isinstance(blob, bytes) and len(blob) > 0

    # Parse the DOCX, find the runs containing each speaker name, and
    # assert color presence/absence.
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(blob))
    primary_runs = []
    moderator_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if "Dr. Mueller" in (run.text or ""):
                primary_runs.append(run)
            if "J. Hsu" in (run.text or ""):
                moderator_runs.append(run)

    assert primary_runs, "Dr. Mueller speaker run not found in DOCX"
    assert moderator_runs, "J. Hsu speaker run not found in DOCX"

    p = primary_runs[0]
    assert p.bold is True, "primary speaker run must be bold"
    color = p.font.color.rgb
    assert color is not None, "primary speaker run must have a color set"
    # Compare via string form because RGBColor equality is finicky across versions.
    assert str(color).upper() == "002855", f"primary color expected #002855, got #{color}"

    m = moderator_runs[0]
    assert m.bold is True, "moderator speaker run must remain bold"
    # Moderator must NOT have navy color — falls through to default (None / theme).
    assert m.font.color.rgb is None or str(m.font.color.rgb).upper() != "002855", (
        "moderator speaker run must NOT be navy"
    )


def test_to_docx_speaker_role_none_falls_back_to_bold_only():
    """Backward compatibility: speaker with role=None still bolds, no color."""
    from app.engines.artifact_transformer import (
        SegmentForExport,
        SessionForExport,
        to_docx,
    )

    seg = SegmentForExport(
        seq=0, start_ms=0, end_ms=1000, text="hello",
        slide_index=0, slide_title="Slide 1",
        speaker_name="Anon", speaker_role=None,
    )
    session = SessionForExport(
        code="t", title="t", presenter=None, duration_sec=None,
        segments=[seg], slides=[],
    )
    blob = to_docx(session)

    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(blob))
    found = False
    for para in doc.paragraphs:
        for run in para.runs:
            if "Anon" in (run.text or ""):
                assert run.bold is True
                assert run.font.color.rgb is None or str(run.font.color.rgb).upper() != "002855"
                found = True
    assert found, "Anon speaker run not found"
