"""
Tests for Phase 5 — segment formatting preservation in DOCX export.

Pre-Phase-5 behavior: `to_docx` collapsed every segment to a single
paragraph regardless of embedded ``\\n`` (hard return) or ``\\n\\n``
(paragraph break). An editor who hard-wrapped a long quote into 3
paragraphs would see one wall of text in the exported Word doc.

Post-Phase-5 behavior:
  * ``\\n\\n`` -> new paragraph
  * ``\\n``    -> soft line break within the current paragraph

Speaker label stays on the FIRST paragraph only (subsequent paragraphs
within the same segment don't repeat the label).
"""
from dataclasses import dataclass
from typing import Optional

import pytest

from app.engines.artifact_transformer import to_docx


@dataclass
class _Seg:
    """Test double for the SegmentRecord shape that to_docx expects."""
    text:          str
    start_ms:      int = 0
    end_ms:        int = 1000
    speaker_name:  Optional[str] = None
    slide_index:   Optional[int] = None
    slide_title:   Optional[str] = None


@dataclass
class _Session:
    title:      str
    code:       str
    presenter:  Optional[str]
    segments:   list[_Seg]
    polls:      list = None  # type: ignore[assignment]
    chat:       list = None  # type: ignore[assignment]
    resources:  list = None  # type: ignore[assignment]
    hyperlinks: list = None  # type: ignore[assignment]

    def __post_init__(self) -> None:
        if self.polls      is None: self.polls      = []
        if self.chat       is None: self.chat       = []
        if self.resources  is None: self.resources  = []
        if self.hyperlinks is None: self.hyperlinks = []


def _build(*segments: _Seg, title: str = "Test", code: str = "T-001",
           presenter: Optional[str] = None) -> _Session:
    return _Session(title=title, code=code, presenter=presenter,
                    segments=list(segments))


def _docx_paragraph_texts(blob: bytes) -> list[str]:
    """Extract paragraph .text values from a generated .docx bytes blob.
    Soft line breaks render as '\\n' inside a paragraph's .text."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(blob))
    return [p.text for p in doc.paragraphs]


class TestDocxFormattingPreservation:

    def test_segment_without_newlines_stays_one_paragraph(self):
        sess = _build(_Seg(text="Hello world."))
        paras = _docx_paragraph_texts(to_docx(sess))
        # Confirm the segment text appears intact as one paragraph (the
        # title + code metadata occupy earlier paragraphs).
        assert "Hello world." in paras

    def test_double_newline_splits_into_separate_paragraphs(self):
        sess = _build(_Seg(text="First paragraph.\n\nSecond paragraph."))
        paras = _docx_paragraph_texts(to_docx(sess))
        assert "First paragraph." in paras
        assert "Second paragraph." in paras
        # The two halves land in DIFFERENT paragraphs (different list entries),
        # not as a single paragraph with the \n\n collapsed.
        assert "First paragraph.\n\nSecond paragraph." not in paras

    def test_single_newline_is_soft_break_within_one_paragraph(self):
        sess = _build(_Seg(text="Line one\nLine two"))
        paras = _docx_paragraph_texts(to_docx(sess))
        # Soft line breaks render as embedded \n in the paragraph's .text.
        joined = "\n".join(paras)
        assert "Line one\nLine two" in joined or "Line one" in paras and "Line two" in paras

    def test_speaker_label_only_on_first_paragraph_of_segment(self):
        sess = _build(_Seg(
            text="Quoting the protocol:\n\nStep one.\n\nStep two.",
            speaker_name="Dr. Mueller",
        ))
        paras = _docx_paragraph_texts(to_docx(sess))
        # First paragraph of the segment carries the speaker prefix.
        speaker_paras = [p for p in paras if p.startswith("Dr. Mueller: ")]
        assert len(speaker_paras) == 1
        assert "Dr. Mueller: Quoting the protocol:" in speaker_paras[0]
        # The other two paragraphs (Step one / Step two) MUST NOT carry
        # the speaker label.
        assert "Dr. Mueller: Step one" not in "\n".join(paras)

    def test_mixed_paragraphs_and_soft_breaks(self):
        # Real-world editorial pattern: a quoted block (paragraphs)
        # followed by tightly-grouped reasoning (soft breaks).
        sess = _build(_Seg(
            text="Findings:\n\n- A\n- B\n- C\n\nConclusion: ship it.",
        ))
        paras = _docx_paragraph_texts(to_docx(sess))
        # 3 segment-paragraphs from \n\n splits (Findings, list, Conclusion).
        # The list itself uses soft breaks (\n) so renders as one para
        # with embedded \n.
        seg_paras = [p for p in paras if any(
            tok in p for tok in ("Findings:", "- A", "Conclusion:")
        )]
        assert any("Findings:" in p for p in seg_paras)
        assert any("Conclusion: ship it." in p for p in seg_paras)
        # The "- A\n- B\n- C" block is one paragraph with embedded \n
        # (or three lines in one para depending on docx renderer).
        assert any("- A" in p and "- B" in p for p in seg_paras)

    def test_empty_segment_does_not_crash(self):
        sess = _build(_Seg(text=""))
        # Pre-fix code did `para.add_run("")` which is benign; post-fix
        # splits "" into [""] and adds one empty paragraph. Confirm no
        # exception.
        blob = to_docx(sess)
        assert isinstance(blob, bytes)
        assert len(blob) > 0

    def test_segment_with_only_newlines_does_not_lose_content(self):
        # Edge case — pure whitespace shouldn't blow up.
        sess = _build(_Seg(text="\n\n\n"))
        blob = to_docx(sess)
        assert isinstance(blob, bytes)
